import urllib.parse
from dotenv import load_dotenv
import os
import json
import asyncio
import logging
import traceback
from pydantic_ai import Agent
from pydantic_ai.mcp import MCPServerSSE, MCPServerStdio
from pydantic_ai import RunContext
import logfire
from docx import Document

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)
logfire.configure()  
logfire.instrument_pydantic_ai() 

# Get environment variables for Coral
runtime = os.getenv("CORAL_ORCHESTRATION_RUNTIME", "devmode")

if runtime == "docker" or runtime == "executable":
    base_url = os.getenv("CORAL_SSE_URL")
    agentID = os.getenv("CORAL_AGENT_ID")
else:
    load_dotenv(override=True)
    base_url = os.getenv("CORAL_SSE_URL")
    agentID = os.getenv("CORAL_AGENT_ID")

# Debug environment variables
logger.info(f"CORAL_SSE_URL: {base_url}")
logger.info(f"CORAL_AGENT_ID: {agentID}")

coral_params = {
    "agentId": agentID,
    "agentDescription": "PowerPoint presentation agent capable of creating comprehensive presentations with tables, charts, shapes, and professional design elements based on user requests."
}

query_string = urllib.parse.urlencode(coral_params)

async def get_mcp_tools(server):
    """Get tools from an MCP server and return formatted tool descriptions."""
    try:
        async with server:
            logger.info(f"Attempting to list tools from server: {server}")
            tools_result = await server.list_tools()
            logger.info(f"Tools result: {tools_result}")
            # Check if tools_result is a list; if so, use it directly
            tools = tools_result if isinstance(tools_result, list) else (tools_result.tools if hasattr(tools_result, 'tools') else [])
            logger.info(f"Retrieved tools: {tools}")
            #commented it out to avoid rate limit error
            # Format tools like get_tools_description
            # def serialize_schema(tool):
            #     # Try different possible schema attributes
            #     if hasattr(tool, 'inputSchema') and tool.inputSchema:
            #         return json.dumps(tool.inputSchema).replace('{', '{{').replace('}', '}}')
            #     elif hasattr(tool, 'parameters_json_schema') and tool.parameters_json_schema:
            #         return json.dumps(tool.parameters_json_schema).replace('{', '{{').replace('}', '}}')
            #     elif hasattr(tool, 'args') and tool.args:
            #         return json.dumps(tool.args).replace('{', '{{').replace('}', '}}')
            #     return "{}"
            
            formatted_tools = "\n".join(
                f"Tool: {tool.name}"
                for tool in tools
            )
            return formatted_tools or "No tools available"
            
    except Exception as e:
        logger.error(f"Error retrieving tools from server {server}: {str(e)}")
        logger.error(traceback.format_exc())
        return "Error retrieving tools"


async def main():
    """Main function to run the PowerPoint agent with Coral integration"""
    try:
        # Setup MCP servers
        CORAL_SERVER_URL = f"{base_url}?{query_string}"
        logger.info(f"Connecting to Coral Server: {CORAL_SERVER_URL}")
        
        # Initialize Coral MCP server (SSE)
        coral_server = MCPServerSSE(
            url=CORAL_SERVER_URL,
            sse_read_timeout=600,
            timeout=600,
        )
        
        # Initialize PowerPoint MCP server (stdio)
        ppt_server = MCPServerStdio(
            command='uvx',
            args=['--from', 'office-powerpoint-mcp-server', 'ppt_mcp_server'],
            timeout=600,
        )
        
        # Get tools from MCP servers before creating agent
        logger.info("Getting tools from MCP servers...")
        coral_tools = await get_mcp_tools(coral_server)
        ppt_tools = await get_mcp_tools(ppt_server)
        
        
        
        # Create system prompt with tool descriptions for coral workflow
        system_prompt = f"""You are a PowerPoint presentation agent interacting with the Coral Server and having PowerPoint creation tools. Your task is to create comprehensive presentations when instructed by other agents.

Follow these steps in order:
1. Call wait_for_mentions from coral tools (timeoutMs: 30000) to receive mentions from other agents.
2. When you receive a mention, keep the thread ID and the sender ID.
3. Analyze the received message to understand what type of presentation is requested.
4. Before creating the presentation, gather all necessary details through counter-questions:
   - First, analyze if the request already contains all necessary details
   - If more details are needed, use send_message to ask follow-up questions about:
     * Number of slides needed (if not specified)
     * Chart/visual requirements (if relevant) - types like bar charts, pie charts, etc.
     * Table requirements (if relevant)
     * Image requirements and their paths/sources (if images are needed)
     * Theme preferences (professional_gradient, gradient or solid)
     * Color scheme preferences (Modern Blue, Corporate Gray, Elegant Green, Warm Red)
     * Transition effects between slides (fade, slide, zoom, etc.)
     * Any reference PowerPoint file to use as a template
     * Any Word document to convert to PowerPoint
     * also the path where it should be saved
   - Only ask questions that are relevant to the specific request
   - Wait for responses using wait_for_mentions after each question
   - After gathering all details, summarize the requirements and ask for final confirmation
   - Only proceed with presentation creation after receiving confirmation

5. Once confirmed, use your PowerPoint tools to create comprehensive presentations with:
   - Clear structure and flow that is well organized and ordered
   - Informative tables showing comparisons or specifications when relevant
   - Relevant charts and diagrams to visualize data when asked for it
   - Professional shapes and design elements
   - Beautiful and consistent formatting
   - Everything properly organized and sequenced
6. When creating presentations, save them in the directory provided by the user.
7. Make sure presentations are both informative and visually appealing with professional design.
8. Think about the content and see if you have executed the instruction to the best of your ability using the PowerPoint tools. Make this your response as "answer" including details about what was created.
9. Use `send_message` from coral tools to send a message in the same thread ID to the sender Id you received the mention from, with content: "answer" and the directory where the presentation was saved.
10. If any error occurs, use `send_message` to send a message in the same thread ID to the sender Id you received the mention from, with content: "error" with details about what went wrong.
11. Always respond back to the sender agent even if you have no answer or error.
12. Repeat the process from step 1.

These are the coral tools available: {coral_tools}

These are the PowerPoint tools available: {ppt_tools}

When creating presentations, ensure they include:
- Professional layouts and design
- Proper structure with clear sections
- Beautiful formatting and consistent styling
- Text optimization for every slide:
  * Appropriate font sizes that are readable but not oversized
  * No overcrowding of text on any slide
  * Use of bullet points and hierarchical structure when appropriate
  * Text properly aligned with other elements on the slide
- Make sure that text is well fitted in the slides without overflow or cramping
- If you have added visuals like charts or graphs then do not include bullet points in it as it will be messy

Remember:
- Only ask for details that are not already provided in the initial request
- Be context-aware when asking questions
- Always get final confirmation before creating the presentation
- Maintain professional communication throughout the process
"""
        
        # Initialize agent with both coral and PowerPoint tools
        agent = Agent(
            model=f"{os.getenv('MODEL_PROVIDER', 'openai')}:{os.getenv('MODEL_NAME', 'gpt-4o-mini')}",
            system_prompt=system_prompt,
            mcp_servers=[coral_server, ppt_server]
        )

        @agent.tool
        def read_word_document(ctx: RunContext, file_path: str) -> str:
            """
            Read a Word document and return its content.
            Args:
                ctx: The run context for the tool
                file_path: The full path to the Word document (.docx file)
            Returns:
                The text content of the Word document
            """
            try:
                doc = Document(file_path)
                content = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        content.append(paragraph.text.strip())
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                content.append(cell.text.strip())
                return "\n\n".join(content)
            except Exception as e:
                return f"Error reading document: {str(e)}"

        logger.info("Coral and PowerPoint MCP Server Connection Established")
        
        # Initialize message history
        message_history = []

        # Run the agent with MCP servers
        async with agent.run_mcp_servers():
            logger.info("=== CONNECTION ESTABLISHED ===")
            
            while True:
                try:
                    logger.info("Starting new agent invocation")
                    
                    # Run the agent to wait for mentions and process them
                    result = await agent.run(
                        "Call wait_for_mentions to wait for presentation creation instructions from other agents.When you are making the slides make sure that the text is well fitted in the slides and use proffesional gradients and colors in the slides",
                        message_history=message_history
                    )
                    
                    logger.info(f"Agent result: {result.output}")
                    
                    # Update message history with new messages
                    message_history.extend(result.new_messages())
                    
                    # Log the current message history size
                    logger.debug(f"Current message history size: {len(message_history)}")
                    
                    logger.info("Completed agent invocation, restarting loop")
                    await asyncio.sleep(1)
                    
                except Exception as e:
                    logger.error(f"Error in agent loop: {str(e)}")
                    logger.error(traceback.format_exc())
                    await asyncio.sleep(5)

    except Exception as e:
        logger.error(f"Error in main setup: {str(e)}")
        logger.error(traceback.format_exc())
        raise

if __name__ == "__main__":
    asyncio.run(main())
